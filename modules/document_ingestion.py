"""
Module: document_ingestion.py
Handles extraction of raw text + structured chunks from:
PDF, DOCX, XLSX, .msg/.eml email, HTML (ACORD-style), TXT

Fixes vs original:
- source_quality field: "rich" | "table_row" | "header" — used by section_weighter
- ACORD boilerplate rows filtered (agency/contact/form-label rows)
- Table rows paired with column headers for context
- Deduplication between table and paragraph extraction
"""

import re
import json
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional
from loguru import logger


@dataclass
class DocumentChunk:
    text: str
    page: Optional[int]       = None
    section: Optional[str]    = None
    chunk_index: int           = 0
    source_type: str           = ""
    source_quality: str        = "rich"   # rich | table_row | header | boilerplate
    table_data: Optional[list] = None
    raw_metadata: dict         = field(default_factory=dict)


@dataclass
class IngestedDocument:
    filename: str
    source_type: str
    chunks: List[DocumentChunk]
    full_text: str
    metadata: dict = field(default_factory=dict)


# ── Boilerplate filter ─────────────────────────────────────────────────────────
_BOILERPLATE_RE = re.compile(
    r"^(AGENCY CUSTOMER ID|APPLICANT INFORMATION|NAME \(Other Named|"
    r"MAILING ADDRESS|GL CODE|SIC\b|NAICS\b|FEIN\b|BUSINESS PHONE|"
    r"WEBSITE ADDRESS|CORPORATION\s+JOINT VENTURE|NOT FOR PROFIT|"
    r"SUBCHAPTER|NO\. OF MEMBERS|CONTACT TYPE|CONTACT NAME|"
    r"PRIMARY HOME BUS CELL|SECONDARY HOME BUS CELL|PHONE # PHONE #|"
    r"Y / N\s+\d|EXPLAIN ALL .YES. RESPONSES|IS THE APPLICANT A SUBSIDIARY|"
    r"DOES THE APPLICANT HAVE ANY|IS A FORMAL SAFETY PROGRAM|"
    r"ANY EXPOSURE TO FLAMMABLES|ANY OTHER INSURANCE WITH THIS|"
    r"ANY POLICY OR COVERAGE DECLINED|ADDITIONAL INTEREST\s*\(|"
    r"INTEREST NAME AND ADDRESS|LEASEBACK REGISTRANT|"
    r"LOSS PAYABLE LIEN AMOUNT|LINE OF BUSINESS POLICY NUMBER|"
    r"Coverage Description \| Limit Amount)",
    re.IGNORECASE,
)

def _is_boilerplate(text: str) -> bool:
    return bool(_BOILERPLATE_RE.match(text.strip()))

def _is_meaningful(text: str, min_chars: int = 20, min_words: int = 3) -> bool:
    t = text.strip()
    if len(t) < min_chars or len(t.split()) < min_words:
        return False
    alpha = sum(1 for c in t if c.isalpha())
    return alpha / max(len(t), 1) >= 0.25


def ingest_document(file_path: str) -> IngestedDocument:
    path = Path(file_path)
    ext  = path.suffix.lower()
    logger.info(f"Ingesting {path.name} ({ext})")
    if ext == ".pdf":        return _ingest_pdf(path)
    elif ext == ".docx":     return _ingest_docx(path)
    elif ext in (".xlsx", ".xls"): return _ingest_excel(path)
    elif ext in (".msg", ".eml"):  return _ingest_email(path)
    elif ext in (".html", ".htm", ".xml"): return _ingest_html_acord(path)
    elif ext == ".txt":      return _ingest_txt(path)
    else: raise ValueError(f"Unsupported file type: {ext}")


def _ingest_pdf(path: Path) -> IngestedDocument:
    import pdfplumber
    chunks, full_parts, idx = [], [], 0
    table_seen = set()

    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            for table in page.extract_tables() or []:
                if not table: continue
                header_row, data_rows = None, table
                if len(table) > 1:
                    candidate = [str(c).strip() if c else "" for c in table[0]]
                    if len([c for c in candidate if c]) >= len(candidate) * 0.5:
                        header_row, data_rows = candidate, table[1:]
                for row in data_rows:
                    cells = [str(c).strip() if c else "" for c in row]
                    if header_row:
                        pairs = [f"{h}: {v}" for h, v in zip(header_row, cells) if h and v]
                        rt = " | ".join(pairs) if pairs else " | ".join(c for c in cells if c)
                    else:
                        rt = " | ".join(c for c in cells if c)
                    if not _is_meaningful(rt, 15, 3) or _is_boilerplate(rt): continue
                    table_seen.add(rt.lower()[:60])
                    chunks.append(DocumentChunk(text=rt, page=page_num,
                        section=_guess_section(rt) or "table",
                        chunk_index=idx, source_type="pdf", source_quality="table_row",
                        table_data=row))
                    idx += 1; full_parts.append(rt)

            for para in _split_paragraphs(page.extract_text() or ""):
                if not _is_meaningful(para) or _is_boilerplate(para): continue
                if para.lower()[:60] in table_seen: continue
                chunks.append(DocumentChunk(text=para, page=page_num,
                    section=_guess_section(para), chunk_index=idx,
                    source_type="pdf", source_quality="rich"))
                idx += 1; full_parts.append(para)

    return IngestedDocument(filename=path.name, source_type="pdf",
        chunks=chunks, full_text="\n".join(full_parts))


def _ingest_docx(path: Path) -> IngestedDocument:
    from docx import Document
    doc = Document(str(path))
    chunks, full_parts, idx = [], [], 0
    current_section = "body"

    for block in doc.element.body:
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph as DocxPara
            para  = DocxPara(block, doc)
            text  = para.text.strip()
            if not text: continue
            style = para.style.name if para.style else ""
            if "Heading" in style:
                current_section = text
                continue
            if not _is_meaningful(text) or _is_boilerplate(text): continue
            chunks.append(DocumentChunk(text=text, section=current_section,
                chunk_index=idx, source_type="docx", source_quality="rich",
                raw_metadata={"style": style}))
            idx += 1; full_parts.append(text)
        elif tag == "tbl":
            from docx.table import Table as DocxTable
            tbl  = DocxTable(block, doc)
            rows = [[c.text.strip() for c in row.cells] for row in tbl.rows]
            header_row, data_rows = None, rows
            if rows and len([c for c in rows[0] if c]) >= len(rows[0]) * 0.5:
                header_row, data_rows = rows[0], rows[1:]
            for row in data_rows:
                if header_row:
                    pairs = [f"{h}: {v}" for h, v in zip(header_row, row) if h and v]
                    rt = " | ".join(pairs) if pairs else " | ".join(c for c in row if c)
                else:
                    rt = " | ".join(c for c in row if c)
                if not _is_meaningful(rt, 15, 3) or _is_boilerplate(rt): continue
                chunks.append(DocumentChunk(text=rt, section=current_section,
                    chunk_index=idx, source_type="docx", source_quality="table_row",
                    table_data=row))
                idx += 1; full_parts.append(rt)

    return IngestedDocument(filename=path.name, source_type="docx",
        chunks=chunks, full_text="\n".join(full_parts))


def _ingest_excel(path: Path) -> IngestedDocument:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), data_only=True)
    chunks, full_parts, idx = [], [], 0
    for sheet_name in wb.sheetnames:
        ws   = wb[sheet_name]
        rows = [[str(c).strip() if c is not None else "" for c in row]
                for row in ws.iter_rows(values_only=True) if any(c is not None for c in row)]
        if not rows: continue
        header_row, data_rows = None, rows
        if len(rows) > 1 and len([c for c in rows[0] if c]) >= 2:
            header_row, data_rows = rows[0], rows[1:]
        for row in data_rows:
            if header_row:
                pairs = [f"{h}: {v}" for h, v in zip(header_row, row) if h and v]
                rt = " | ".join(pairs) if pairs else " | ".join(c for c in row if c)
            else:
                rt = " | ".join(c for c in row if c)
            if not _is_meaningful(rt, 15, 3) or _is_boilerplate(rt): continue
            chunks.append(DocumentChunk(text=rt, section=f"sheet:{sheet_name}",
                chunk_index=idx, source_type="xlsx", source_quality="table_row",
                table_data=row, raw_metadata={"sheet": sheet_name}))
            idx += 1; full_parts.append(rt)
    return IngestedDocument(filename=path.name, source_type="xlsx",
        chunks=chunks, full_text="\n".join(full_parts))


def _ingest_email(path: Path) -> IngestedDocument:
    chunks, full_parts, idx = [], [], 0
    if path.suffix.lower() == ".msg":
        try:
            import extract_msg
            msg  = extract_msg.Message(str(path))
            body = msg.body or ""
            meta = {"subject": msg.subject or "", "sender": msg.sender or ""}
        except Exception as e:
            logger.warning(f"extract_msg failed: {e}")
            body, meta = path.read_text(errors="ignore"), {}
    else:
        import email as _email
        msg  = _email.message_from_bytes(path.read_bytes())
        body = ""
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    body += part.get_payload(decode=True).decode(errors="ignore")
        else:
            body = msg.get_payload(decode=True).decode(errors="ignore")
        meta = {"subject": msg.get("Subject",""), "sender": msg.get("From","")}

    for para in _split_paragraphs(body):
        if not _is_meaningful(para): continue
        chunks.append(DocumentChunk(text=para, section="email_body",
            chunk_index=idx, source_type="email", source_quality="rich"))
        idx += 1; full_parts.append(para)
    return IngestedDocument(filename=path.name, source_type="email",
        chunks=chunks, full_text="\n".join(full_parts), metadata=meta)


def _ingest_html_acord(path: Path) -> IngestedDocument:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(path.read_text(errors="ignore"), "lxml")
    chunks, full_parts, idx = [], [], 0
    for table in soup.find_all("table"):
        header_cells = []
        for tr in table.find_all("tr"):
            cells = [td.get_text(strip=True) for td in tr.find_all(["td","th"])]
            if not any(cells): continue
            if not header_cells: header_cells = cells; continue
            if header_cells:
                pairs = [f"{h}: {v}" for h, v in zip(header_cells, cells) if h and v]
                rt = " | ".join(pairs) if pairs else " | ".join(c for c in cells if c)
            else:
                rt = " | ".join(c for c in cells if c)
            if not _is_meaningful(rt, 15, 3) or _is_boilerplate(rt): continue
            chunks.append(DocumentChunk(text=rt, section="acord_table",
                chunk_index=idx, source_type="acord", source_quality="table_row",
                table_data=cells))
            idx += 1; full_parts.append(rt)
    for tag in soup.find_all(["p","div","li"]):
        text = tag.get_text(strip=True)
        if not _is_meaningful(text) or _is_boilerplate(text): continue
        chunks.append(DocumentChunk(text=text, section=_guess_section(text),
            chunk_index=idx, source_type="acord", source_quality="rich"))
        idx += 1; full_parts.append(text)
    return IngestedDocument(filename=path.name, source_type="acord",
        chunks=chunks, full_text="\n".join(full_parts))


def _ingest_txt(path: Path) -> IngestedDocument:
    text   = path.read_text(errors="ignore")
    chunks = []
    for i, para in enumerate(_split_paragraphs(text)):
        if not _is_meaningful(para): continue
        chunks.append(DocumentChunk(text=para, section=_guess_section(para),
            chunk_index=i, source_type="txt", source_quality="rich"))
    return IngestedDocument(filename=path.name, source_type="txt",
        chunks=chunks, full_text=text)


def _split_paragraphs(text: str) -> List[str]:
    result = []
    for p in re.split(r"\n{2,}", text):
        p = p.strip()
        if not p: continue
        if len(p) > 800:
            result.extend(s.strip() for s in p.split("\n") if s.strip())
        else:
            result.append(p)
    return result


def _table_to_text(table: list) -> str:
    return "\n".join(" | ".join(str(c) if c else "" for c in row) for row in table if row)


SECTION_KEYWORDS = {
    "risk description":  ["risk description","property description","risk details","insured premises"],
    "prior claims":      ["prior claims","loss history","claims history","previous claims","loss runs"],
    "operations":        ["business operations","operations","activities","occupancy","hours of operation"],
    "financial":         ["financial","balance sheet","revenue","premium","payroll","bankruptcy"],
    "safety":            ["safety program","osha","safety manual","ppe","training"],
    "general conditions":["general conditions","exclusions","limitations","endorsements"],
    "signature":         ["signature","authorized","date signed"],
}

def _guess_section(text: str) -> str:
    lower = text.lower()
    for section, kws in SECTION_KEYWORDS.items():
        if any(kw in lower for kw in kws):
            return section
    return "body"


def get_context_for_display(ingested_doc, chunk_index: int,
                             keyword: str, window: int = 1) -> str:
    chunks = ingested_doc.chunks
    if not chunks or chunk_index >= len(chunks):
        return ""
    start    = max(0, chunk_index - window)
    end      = min(len(chunks), chunk_index + window + 1)
    combined = " ".join(c.text for c in chunks[start:end])
    if keyword:
        kw_lower  = keyword.lower()
        sentences = re.split(r"(?<=[.!?])\s+", combined)
        marked, found = [], False
        for sent in sentences:
            if not found and kw_lower in sent.lower():
                marked.append(f"→ {sent}"); found = True
            else:
                marked.append(sent)
        combined = " ".join(marked)
    return combined.strip()

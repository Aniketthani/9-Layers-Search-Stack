"""
generate_samples.py
Creates sample submission documents in multiple formats for testing.
Run once: python generate_samples.py
"""

import json
from pathlib import Path

OUTPUT_DIR = Path("data/sample_documents")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ── Sample 1: DOCX submission with multiple adverse flags ─────────────────────

def create_sample_docx():
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Commercial Property Insurance Submission", 0)

    doc.add_heading("Insured Information", 1)
    doc.add_paragraph("Insured Name: Riverside Industrial Holdings LLC")
    doc.add_paragraph("Location: 4420 Harbor Drive, Port Granville, TX 77001")
    doc.add_paragraph("Business Type: Chemical Storage and Distribution")

    doc.add_heading("Risk Description", 1)
    doc.add_paragraph(
        "The insured operates a chemical storage facility adjacent to the Port Granville "
        "industrial corridor. The property has a history of environmental contamination "
        "relating to underground storage tanks removed in 2018. Soil remediation was "
        "partially completed; groundwater contamination has been confirmed by the EPA "
        "and an enforcement action is currently active."
    )

    doc.add_heading("Prior Claims and Loss History", 1)
    doc.add_paragraph(
        "The insured was involved in prior litigation in 2019 regarding a chemical spill "
        "on the adjacent parcel. The matter was settled out of court for an undisclosed sum. "
        "A class action lawsuit was filed in 2021 by neighboring residents citing health impacts; "
        "this matter is currently unresolved."
    )

    doc.add_heading("Financial Information", 1)
    doc.add_paragraph(
        "The company filed for Chapter 11 bankruptcy protection in Q3 2022 and emerged "
        "from restructuring in Q1 2023. A debt restructuring agreement is in place with "
        "primary lenders."
    )

    doc.add_heading("Property Details", 1)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Feature"
    hdr[1].text = "Details"
    hdr[2].text = "Notes"
    rows_data = [
        ("Construction", "Pre-1980 steel frame", "May contain asbestos insulation"),
        ("Roof", "Original flat membrane", "No known flood damage"),
        ("Sprinkler system", "Partial coverage only", "Fire suppression failure noted in 2020 inspection"),
        ("Occupancy", "Chemical storage — flammable", "Permit violation cited by fire marshal 2022"),
    ]
    for feat, detail, note in rows_data:
        row = t.add_row().cells
        row[0].text = feat
        row[1].text = detail
        row[2].text = note

    doc.add_heading("General Conditions", 1)
    doc.add_paragraph(
        "This submission is subject to standard exclusions. Environmental liability, "
        "pollution, and prior litigation exclusions apply as standard. "
        "No coverage is provided for known asbestos remediation costs."
    )

    out = OUTPUT_DIR / "sample_submission_high_risk.docx"
    doc.save(str(out))
    print(f"Created: {out}")


# ── Sample 2: PDF (via reportlab or plain text fallback) ──────────────────────

def create_sample_txt_as_pdf_fallback():
    """Creates a .txt file that represents a clean low-risk submission."""
    content = """COMMERCIAL PROPERTY SUBMISSION FORM
Submitted by: Thompson Valley Real Estate Group

PROPERTY DESCRIPTION
Location: 88 Clearwater Boulevard, Springfield, OH 44101
Use: Mixed-use retail and office, 3 stories
Year Built: 2009
Construction: Steel frame, Class A

LOSS HISTORY
No prior claims in the past 5 years.
No history of litigation related to this property.
No known environmental issues or contamination.
Building has passed all fire code inspections with no violations cited.

FINANCIAL OVERVIEW
The insured entity is financially stable with no bankruptcy filings.
No outstanding creditor claims or foreclosure actions.

OCCUPANCY
Ground floor: Retail — general merchandise
Upper floors: Professional offices — law firm and accounting practice
No hazardous materials storage. No chemical operations on premises.
Sprinkler system fully operational. Last certified: March 2024.

ACORD SUPPLEMENTAL DATA
Form: ACORD 140
Policy Type: Commercial Package
Requested Limits: $5,000,000 per occurrence
Deductible: $25,000
"""
    out = OUTPUT_DIR / "sample_submission_low_risk.txt"
    out.write_text(content)
    print(f"Created: {out}")


# ── Sample 3: Excel with ACORD-style tabular data ─────────────────────────────

def create_sample_excel():
    import openpyxl
    from openpyxl.styles import Font, PatternFill

    wb = openpyxl.Workbook()

    # Sheet 1: Property details
    ws1 = wb.active
    ws1.title = "Property Details"
    ws1.append(["Field", "Value", "Underwriter Notes"])
    ws1.append(["Insured Name", "Harborview Chemical Corp", ""])
    ws1.append(["Address", "2200 Industrial Pkwy, Newark NJ", ""])
    ws1.append(["SIC Code", "2890 - Industrial Chemicals", "High risk SIC"])
    ws1.append(["Year Built", "1967", "Pre-regulatory era construction"])
    ws1.append(["Sq Footage", "180,000", ""])
    ws1.append(["Construction", "Masonry/Steel", "Possible asbestos present"])
    ws1.append(["Occupancy", "Chemical storage and processing", "Flammable storage on site"])
    ws1.append(["Annual Revenue", "$12,400,000", ""])

    # Sheet 2: Loss history with adverse entries
    ws2 = wb.create_sheet("Loss History")
    ws2.append(["Year", "Loss Type", "Amount Paid", "Description", "Status"])
    ws2.append(["2021", "Environmental", "$450,000", "Soil contamination cleanup — UST removal", "Closed"])
    ws2.append(["2022", "Liability", "$1,200,000", "Prior litigation — employee injury settlement", "Closed"])
    ws2.append(["2023", "Regulatory", "N/A", "EPA citation for groundwater contamination — enforcement action active", "Open"])
    ws2.append(["2023", "Fire", "$85,000", "Partial fire damage — sprinkler failure in storage bay 3", "Closed"])
    ws2.append(["2024", "Liability", "Pending", "Class action lawsuit filed — pollution claims from neighboring properties", "Open"])

    # Sheet 3: Financial indicators
    ws3 = wb.create_sheet("Financial")
    ws3.append(["Indicator", "Status", "Notes"])
    ws3.append(["Bankruptcy history", "Yes", "Chapter 11 filed 2020, emerged 2021"])
    ws3.append(["Active liens", "Yes", "2 creditor liens on property"])
    ws3.append(["Credit rating", "CCC+", "Below investment grade"])
    ws3.append(["Debt restructuring", "In progress", "Creditor agreement signed Jan 2024"])
    ws3.append(["Foreclosure risk", "Moderate", "Flagged by lender Q4 2023"])

    out = OUTPUT_DIR / "sample_submission_acord_excel.xlsx"
    wb.save(str(out))
    print(f"Created: {out}")


# ── Sample 4: HTML ACORD-style form ───────────────────────────────────────────

def create_sample_html_acord():
    html = """<!DOCTYPE html>
<html>
<head><title>ACORD 125 - Commercial Insurance Application</title></head>
<body>
<h1>ACORD 125 — Commercial Insurance Application</h1>

<h2>Applicant Information</h2>
<table border="1">
  <tr><th>Field</th><th>Response</th></tr>
  <tr><td>Named Insured</td><td>Delta Processing Industries Inc.</td></tr>
  <tr><td>Business Type</td><td>Waste Processing and Recycling</td></tr>
  <tr><td>Years in Business</td><td>14</td></tr>
  <tr><td>FEIN</td><td>47-XXXXXXX</td></tr>
</table>

<h2>Risk Description</h2>
<p>Delta Processing operates a waste processing facility handling both municipal solid waste 
and industrial byproducts. The facility processes approximately 200 tons per day. 
Hazardous waste streams are handled under EPA permit #TX-HW-2019-0041.</p>

<p>A prior regulatory violation was issued in 2022 for improper storage of chemical waste. 
The citation was resolved with a $45,000 fine imposed by the Texas Commission on Environmental Quality.</p>

<h2>Prior Litigation and Claims</h2>
<table border="1">
  <tr><th>Year</th><th>Type</th><th>Description</th><th>Resolution</th></tr>
  <tr><td>2021</td><td>Environmental</td><td>Lawsuit filed alleging groundwater contamination by adjacent landowner</td><td>Settlement for $320,000</td></tr>
  <tr><td>2022</td><td>Regulatory</td><td>OSHA violation — lack of protective equipment in processing area</td><td>$18,000 fine, corrective action completed</td></tr>
  <tr><td>2023</td><td>Liability</td><td>Class action filed by neighborhood association citing pollution</td><td>Pending — trial set for Q3 2025</td></tr>
</table>

<h2>Property Information</h2>
<table border="1">
  <tr><th>Feature</th><th>Detail</th></tr>
  <tr><td>Location</td><td>Industrial Zone 4, Houston TX</td></tr>
  <tr><td>Building Age</td><td>1972 — possible asbestos in roofing material</td></tr>
  <tr><td>Storage</td><td>Flammable storage — Class II liquids</td></tr>
  <tr><td>Fire Protection</td><td>Partial sprinkler — fire code violation noted 2023</td></tr>
</table>

<h2>General Exclusions</h2>
<p>Standard pollution exclusion applies. Prior litigation exclusion applies. 
Environmental liability for known contamination is excluded from coverage.</p>
</body>
</html>
"""
    out = OUTPUT_DIR / "sample_acord_form.html"
    out.write_text(html)
    print(f"Created: {out}")


# ── Sample 5: Email submission ────────────────────────────────────────────────

def create_sample_email():
    eml_content = """From: broker@insurancebrokers.com
To: underwriting@insuranceco.com
Subject: New Submission - Harborview Metals LLC - Commercial Property

Dear Underwriting Team,

Please find below the submission details for our client Harborview Metals LLC.

INSURED: Harborview Metals LLC
LOCATION: 550 Steel Works Drive, Pittsburgh PA 15201
COVERAGE REQUESTED: Commercial Property + General Liability

BACKGROUND:
The insured is a steel fabrication and metal processing company operating 
since 1985. The facility is approximately 95,000 square feet on a brownfield 
site that was previously used for coal processing.

KNOWN ISSUES:
- The property has documented soil contamination from prior industrial use. 
  Environmental remediation was completed in 2019 but groundwater monitoring 
  is ongoing per EPA requirements.
- The insured has no prior litigation history related to the current operations.
- A structural assessment in 2023 identified foundation settling in the 
  northwest corner of the main building. Remedial work is scheduled for Q2 2024.
- The facility was cited for an OSHA violation in 2022 relating to crane 
  safety procedures. This has been fully remediated.

FINANCIAL:
The insured is financially stable. No bankruptcy or foreclosure history.
Annual revenue approximately $28 million.

Please let me know if you require additional information.

Best regards,
James Morton
Senior Commercial Broker
Morton & Associates Insurance
"""
    out = OUTPUT_DIR / "sample_email_submission.eml"
    out.write_text(eml_content)
    print(f"Created: {out}")


if __name__ == "__main__":
    print("Generating sample documents...")
    create_sample_docx()
    create_sample_txt_as_pdf_fallback()
    create_sample_excel()
    create_sample_html_acord()
    create_sample_email()
    print("\nAll sample documents created in data/sample_documents/")
